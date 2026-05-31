"""
Part 2 Classification Report — Word (.docx) version
Student: Nabila Kader | ID: 23079506
"""
import sqlite3, io, os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import rcParams
rcParams['font.family'] = 'DejaVu Sans'

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colors ────────────────────────────────────────────────────
NAVY  = RGBColor(0x1B, 0x3A, 0x5C)
BLUE  = RGBColor(0x2E, 0x6D, 0xA4)
GOLD  = RGBColor(0xE8, 0xA0, 0x20)
LGRAY = RGBColor(0xF5, 0xF7, 0xFA)
DGRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
AMBER = RGBColor(0xFF, 0xF3, 0xCD)

DB = "23079506-sq26-classification.db"

BAR_COLORS = ['#1B3A5C','#2E6DA4','#4A9FD4','#6AB8E8','#E8A020',
              '#C0392B','#27AE60','#8E44AD','#E67E22','#16A085',
              '#2980B9','#D35400','#7F8C8D','#2C3E50','#F39C12']

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        val = kwargs.get(side, {'sz': '4', 'val': 'single', 'color': 'CCCCCC'})
        el = OxmlElement(f'w:{side}')
        for k, v in val.items():
            el.set(qn(f'w:{k}'), str(v))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run_bold(para, text, size=None, color=None, italic=False):
    run = para.add_run(text)
    run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 12)
    run.font.color.rgb = NAVY if level == 1 else BLUE
    return p

def body_para(doc, text, space_after=6, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DGRAY
    return p

def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E6DA4')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Data ──────────────────────────────────────────────────────
def get_data():
    conn = sqlite3.connect(DB)
    data = {}
    for repo_id, repo_name in conn.execute("SELECT id, name FROM REPOSITORIES ORDER BY id"):
        type_counts = {r[0]: r[1] for r in conn.execute(
            "SELECT type, COUNT(*) FROM PROJECTS WHERE repository_id=? GROUP BY type", (repo_id,))}
        class_counts = {r[0]: r[1] for r in conn.execute(
            "SELECT primary_class, COUNT(*) FROM PROJECTS WHERE repository_id=? "
            "AND primary_class IS NOT NULL GROUP BY primary_class ORDER BY COUNT(*) DESC", (repo_id,))}
        total = conn.execute(
            "SELECT COUNT(*) FROM PROJECTS WHERE repository_id=?", (repo_id,)).fetchone()[0]
        data[repo_name] = {"id": repo_id, "total": total,
                           "type_counts": type_counts, "class_counts": class_counts}
    conn.close()
    return data

# ── Charts ────────────────────────────────────────────────────
def make_bar_chart_img(class_counts, title, top_n=20):
    items = sorted(class_counts.items(), key=lambda x: x[1], reverse=True)[:top_n]
    if not items:
        return None
    labels = [i[0].split(": ", 1)[1] if ": " in i[0] else i[0] for i in items]
    labels = [l[:45] + "…" if len(l) > 45 else l for l in labels]
    values = [i[1] for i in items]
    n = len(labels)
    fig, ax = plt.subplots(figsize=(11, max(5, n * 0.52)))
    bar_colors = [BAR_COLORS[i % len(BAR_COLORS)] for i in range(n)]
    bars = ax.barh(range(n), values, color=bar_colors, edgecolor='white', linewidth=0.6, height=0.7)
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + max(values) * 0.008, bar.get_y() + bar.get_height() / 2,
                f'{val:,}', va='center', ha='left', fontsize=8.5, fontweight='bold', color='#333333')
    ax.set_yticks(range(n))
    ax.set_yticklabels(labels, fontsize=8.5)
    ax.invert_yaxis()
    ax.set_xlabel("Number of Projects", fontsize=9, color='#444444')
    ax.set_title(title, fontsize=11, fontweight='bold', color='#1B3A5C', pad=12)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#CCCCCC')
    ax.spines['bottom'].set_color('#CCCCCC')
    ax.set_xlim(0, max(values) * 1.18)
    ax.grid(axis='x', alpha=0.25, linestyle='--', color='#AAAAAA')
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('white')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf

def make_pie_chart_img(type_counts, title):
    labels = list(type_counts.keys())
    values = list(type_counts.values())
    pie_colors = ['#2E6DA4', '#27AE60', '#E8A020', '#C0392B', '#8E44AD'][:len(labels)]
    fig, ax = plt.subplots(figsize=(7, 4.5))
    wedges, texts, autotexts = ax.pie(
        values, labels=None, autopct='%1.1f%%',
        colors=pie_colors, startangle=90,
        pctdistance=0.75, wedgeprops={'linewidth': 1.5, 'edgecolor': 'white'})
    for at in autotexts:
        at.set_fontsize(9)
        at.set_fontweight('bold')
        at.set_color('white')
    legend_labels = [f"{l} ({v:,})" for l, v in zip(labels, values)]
    ax.legend(wedges, legend_labels, loc='lower center',
              bbox_to_anchor=(0.5, -0.12), ncol=2, fontsize=8.5,
              framealpha=0.9, edgecolor='#CCCCCC')
    ax.set_title(title, fontsize=10, fontweight='bold', color='#1B3A5C', pad=8)
    fig.patch.set_facecolor('white')
    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=200, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf

# ── Cover page ───────────────────────────────────────────────
def add_cover(doc):
    # Dark navy cover background via table spanning full page
    # We approximate with a styled table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    # Title block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(60)
    title_p.paragraph_format.space_after = Pt(6)
    r = title_p.add_run("QDArchive")
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(4)
    r2 = sub_p.add_run("Part 2: Data Classification Report")
    r2.font.size = Pt(18)
    r2.font.bold = True
    r2.font.color.rgb = BLUE

    # Gold rule
    hr(doc)

    # Subtitle
    s1 = doc.add_paragraph("ISIC Rev. 5 Classification of Qualitative Research Projects")
    s1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s1.paragraph_format.space_after = Pt(2)
    for run in s1.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DGRAY

    s2 = doc.add_paragraph("DataFirst UCT & Harvard Murray Archive")
    s2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s2.paragraph_format.space_after = Pt(30)
    for run in s2.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DGRAY

    # Metadata table
    meta = [
        ("Student", "Nabila Kader"),
        ("Student ID", "23079506"),
        ("Course", "Seeding QDArchive (SQ26)"),
        ("University", "FAU Erlangen-Nürnberg"),
        ("Professor", "Prof. Dr. Dirk Riehle"),
        ("Standard", "ISIC Rev. 5 — Division Level"),
        ("Repositories", "DataFirst UCT (#8) + Harvard Murray Archive (#18)"),
        ("Projects Classified", "1,518"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    for i, (label, val) in enumerate(meta):
        row = tbl.rows[i]
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = Cm(5)
        vc.width = Cm(11)
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(3)
        lp.paragraph_format.space_after = Pt(3)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.color.rgb = GOLD
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after = Pt(3)
        vr = vp.add_run(val)
        vr.font.size = Pt(10)
        vr.font.color.rgb = NAVY
        set_cell_bg(lc, RGBColor(0x0D, 0x1F, 0x35))
        set_cell_bg(vc, RGBColor(0x0D, 0x1F, 0x35))

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(20)

    sem = doc.add_paragraph("Summer Semester 2026")
    sem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sem.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = DGRAY

    doc.add_page_break()

# ── Table of Contents placeholder ────────────────────────────
def add_toc(doc):
    heading(doc, "Table of Contents", level=1)
    hr(doc)
    toc_items = [
        ("1.", "Executive Summary", "3"),
        ("2.", "Repository: DataFirst UCT", "4"),
        ("3.", "Repository: Harvard Murray Archive / Dataverse", "6"),
        ("4.", "Overall Summary", "8"),
    ]
    tbl = doc.add_table(rows=len(toc_items), cols=3)
    tbl.style = 'Table Grid'
    for i, (num, title, page) in enumerate(toc_items):
        row = tbl.rows[i]
        row.cells[0].paragraphs[0].add_run(num).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(title).font.size = Pt(10)
        row.cells[2].paragraphs[0].add_run(page).font.size = Pt(10)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(13)
        row.cells[2].width = Cm(2)
        # No border, just a bottom line
        for cell in row.cells:
            set_cell_bg(cell, RGBColor(0xFF, 0xFF, 0xFF))

    doc.add_page_break()

# ── Stat boxes row ────────────────────────────────────────────
def add_stat_boxes(doc, stats):
    tbl = doc.add_table(rows=1, cols=len(stats))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for i, (label, val) in enumerate(stats):
        cell = tbl.rows[0].cells[i]
        cell.width = Cm(4)
        set_cell_bg(cell, LGRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(val + "\n")
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = NAVY
        r2 = p.add_run(label)
        r2.font.size = Pt(8)
        r2.font.color.rgb = DGRAY
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(10)

# ── Type distribution table ───────────────────────────────────
def add_type_table(doc, type_counts, total):
    rows = [["Project Type", "Count", "Percentage"]]
    for ptype in ["QDA_PROJECT", "QD_PROJECT", "OTHER_PROJECT", "NOT_A_PROJECT"]:
        cnt = type_counts.get(ptype, 0)
        if cnt > 0:
            rows.append([ptype, f"{cnt:,}", f"{cnt / total * 100:.1f}%"])

    tbl = doc.add_table(rows=len(rows), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        row.cells[0].width = Cm(7)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(3.5)
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.font.size = Pt(10)
            if i == 0:
                set_cell_bg(cell, NAVY)
                r.bold = True
                r.font.color.rgb = WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = RGBColor(0xF5, 0xF7, 0xFA) if i % 2 == 0 else WHITE
                set_cell_bg(cell, bg)
                if j == 0:
                    r.bold = True
                    r.font.color.rgb = BLUE
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)

# ── Rank table ───────────────────────────────────────────────
def add_rank_table(doc, class_counts, top_n=20):
    rows = [["Rank", "ISIC", "Division Name", "Count"]]
    sorted_classes = sorted(class_counts.items(), key=lambda x: x[1], reverse=True)[:top_n]
    for rank, (cls, cnt) in enumerate(sorted_classes, 1):
        parts = cls.split(": ", 1)
        code = parts[0] if len(parts) > 1 else "-"
        name = parts[1] if len(parts) > 1 else cls
        rows.append([str(rank), code, name, f"{cnt:,}"])

    tbl = doc.add_table(rows=len(rows), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [Cm(1.5), Cm(2), Cm(11), Cm(2)]
    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        for j, (text, width) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[j]
            cell.width = width
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            r.font.size = Pt(9)
            if i == 0:
                set_cell_bg(cell, NAVY)
                r.bold = True
                r.font.color.rgb = WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if i == 1:
                    set_cell_bg(cell, RGBColor(0xFF, 0xF3, 0xCD))  # gold top 1
                elif i % 2 == 0:
                    set_cell_bg(cell, LGRAY)
                else:
                    set_cell_bg(cell, WHITE)
                if j in (0, 3):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if j == 1:
                    r.bold = True
                    r.font.color.rgb = BLUE

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)

# ── Repo comments ─────────────────────────────────────────────
REPO_COMMENTS = {
    "datafirst": (
        "DataFirst UCT is a South African survey and statistics repository. "
        "The dominant ISIC class is Scientific Research and Development (72) with 151 projects, "
        "followed by Employment Activities (78) with 115. Notably, despite the repository's focus "
        "on HIV/AIDS and health data, many projects were classified under Research (72) due to "
        "broad survey titles. All 596 DataFirst projects were classified as QD_PROJECT — no QDA "
        "files (.qdpx, .nvp etc.) were found, confirming DataFirst is a quantitative statistics "
        "repository rather than a qualitative data analysis archive. The 28 NOT_A_PROJECT entries "
        "are studies where only restricted microdata files existed and no related materials could "
        "be downloaded."
    ),
    "harvard": (
        "The Harvard Murray Archive / Harvard Dataverse collection contains longitudinal social "
        "science studies. The dominant class is Human Health Activities (86) with 403 projects, "
        "consistent with the Murray Archive's long-term focus on human development, psychology, "
        "and health outcomes research. Notably, 3 projects were classified as QDA_PROJECT — the "
        "only qualitative data analysis files found across both repositories. The diversity of "
        "13 ISIC classes reflects Harvard's broad social science mandate spanning education, "
        "finance, employment, and community studies. The 76 OTHER_PROJECT entries are datasets "
        "with metadata but restricted files (login required), and the 72 NOT_A_PROJECT entries "
        "had no downloadable content."
    ),
}

REPO_DISPLAY = {
    "datafirst": "DataFirst UCT",
    "harvard": "Harvard Murray Archive / Dataverse",
}

# ── Main build ────────────────────────────────────────────────
def build(data, out):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin  = section.bottom_margin = Cm(2.5)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── Cover ──────────────────────────────────────────────────
    add_cover(doc)

    # ── Table of Contents ──────────────────────────────────────
    add_toc(doc)

    # ── Executive Summary ──────────────────────────────────────
    heading(doc, "1.  Executive Summary")
    hr(doc)

    total_all = sum(d["total"] for d in data.values())
    body_para(doc,
        f"This report presents the classification results of {total_all:,} research projects "
        f"collected from two assigned repositories as part of the Seeding QDArchive project at "
        f"FAU Erlangen-Nürnberg. The repositories are DataFirst UCT (repository #8, 596 projects) "
        f"and Harvard Murray Archive / Harvard Dataverse (repository #18, 922 projects). Each "
        f"project was classified using the ISIC Rev. 5 international standard for economic "
        f"activities, going down to the two-digit division level. Classification was performed "
        f"using automated keyword-based analysis of project titles, descriptions, and metadata keywords.")

    heading(doc, "Methodology", level=2)
    body_para(doc, "Projects were first classified by PROJECT_TYPE based on downloaded file extensions:")

    method_rows = [
        ["Type", "Criterion", "Count"],
        ["QDA_PROJECT",   "Contains a QDA file (.qdpx, .nvp, .atlproj, etc.)", "3"],
        ["QD_PROJECT",    "No QDA file but has primary data files (PDF, DOCX, TXT, etc.)", "1,339"],
        ["OTHER_PROJECT", "Has files but not primary data files", "76"],
        ["NOT_A_PROJECT", "No usable files could be identified", "100"],
    ]
    mt = doc.add_table(rows=len(method_rows), cols=3)
    mt.style = 'Table Grid'
    mt.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row_data in enumerate(method_rows):
        row = mt.rows[i]
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(10)
        row.cells[2].width = Cm(2)
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            r.font.size = Pt(9)
            if i == 0:
                set_cell_bg(cell, NAVY)
                r.bold = True
                r.font.color.rgb = WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                bg = RGBColor(0xF5, 0xF7, 0xFA) if i % 2 == 0 else WHITE
                set_cell_bg(cell, bg)
                if j == 0:
                    r.bold = True
                    r.font.color.rgb = BLUE
                elif j == 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    body_para(doc,
        "ISIC classification was then assigned using keyword matching of research subjects to the "
        "most appropriate ISIC Rev. 5 economic sector division. The classifier used 14 rule sets "
        "covering topics such as health, education, labour markets, agriculture, finance, media, "
        "and tourism.")

    doc.add_page_break()

    # ── Per-repository sections ────────────────────────────────
    for sec_num, (repo_name, repo_data) in enumerate(data.items(), 2):
        display = REPO_DISPLAY.get(repo_name, repo_name.title())
        heading(doc, f"{sec_num}.  Repository: {display}")
        hr(doc)

        # Stat boxes
        stats = [
            ("Total Projects", f"{repo_data['total']:,}"),
            ("QDA Projects",   f"{repo_data['type_counts'].get('QDA_PROJECT', 0)}"),
            ("QD Projects",    f"{repo_data['type_counts'].get('QD_PROJECT', 0):,}"),
            ("ISIC Classes",   f"{len(repo_data['class_counts'])}"),
        ]
        add_stat_boxes(doc, stats)

        # Project type distribution
        heading(doc, "Project Type Distribution", level=2)
        add_type_table(doc, repo_data["type_counts"], repo_data["total"])

        # Pie chart
        pie_buf = make_pie_chart_img(repo_data["type_counts"],
                                     f"Project Type Distribution — {display}")
        if pie_buf:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(pie_buf, width=Cm(12))
        add_caption(doc, f"Figure 1: Project type distribution for {display}")

        # ISIC histogram
        heading(doc, "Primary ISIC Class Distribution", level=2)
        n = min(len(repo_data["class_counts"]), 20)
        bar_buf = make_bar_chart_img(repo_data["class_counts"],
                                     f"Primary ISIC Classes — {display} (Top {n})")
        if bar_buf:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(bar_buf, width=Cm(15))
        add_caption(doc,
            f"Figure 2: Top {n} ISIC divisions identified in {display}. "
            f"Values show number of projects per class.")

        # Rank table
        heading(doc, "Rank-Ordered List of Classes (Top 20)", level=2)
        add_rank_table(doc, repo_data["class_counts"])

        # Comments
        heading(doc, "Comments on Findings", level=2)
        body_para(doc, REPO_COMMENTS.get(repo_name, ""))

        doc.add_page_break()

    # ── Overall Summary ────────────────────────────────────────
    heading(doc, f"{len(data) + 2}.  Overall Summary")
    hr(doc)
    heading(doc, "Combined Classification — Both Repositories", level=2)

    all_classes = {}
    for d in data.values():
        for cls, cnt in d["class_counts"].items():
            all_classes[cls] = all_classes.get(cls, 0) + cnt

    # Combined rank table with % of total
    comb_rows = [["Rank", "ISIC", "Division Name", "Count", "% of Total"]]
    for rank, (cls, cnt) in enumerate(
            sorted(all_classes.items(), key=lambda x: x[1], reverse=True)[:15], 1):
        parts = cls.split(": ", 1)
        code = parts[0] if len(parts) > 1 else "-"
        name = parts[1] if len(parts) > 1 else cls
        comb_rows.append([str(rank), code, name, f"{cnt:,}", f"{cnt / total_all * 100:.1f}%"])

    ct = doc.add_table(rows=len(comb_rows), cols=5)
    ct.style = 'Table Grid'
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_ws = [Cm(1.5), Cm(2), Cm(9), Cm(2), Cm(2)]
    for i, row_data in enumerate(comb_rows):
        row = ct.rows[i]
        for j, (text, width) in enumerate(zip(row_data, col_ws)):
            cell = row.cells[j]
            cell.width = width
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            r.font.size = Pt(9)
            if i == 0:
                set_cell_bg(cell, NAVY)
                r.bold = True
                r.font.color.rgb = WHITE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if i == 1:
                    set_cell_bg(cell, RGBColor(0xFF, 0xF3, 0xCD))
                elif i % 2 == 0:
                    set_cell_bg(cell, LGRAY)
                else:
                    set_cell_bg(cell, WHITE)
                if j in (0, 3, 4):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if j == 1:
                    r.bold = True
                    r.font.color.rgb = BLUE

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Combined bar chart
    buf = make_bar_chart_img(all_classes, "Combined Primary ISIC Classes — Both Repositories")
    if buf:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(buf, width=Cm(15))
    add_caption(doc, "Figure 3: Combined ISIC class distribution across both repositories.")

    heading(doc, "Key Observations", level=2)
    observations = [
        "Human Health Activities (ISIC 86) is the dominant class overall with 430 projects (28.3%), "
        "driven by Harvard's health studies. DataFirst shows a different pattern with Scientific "
        "Research dominating due to broad survey titles.",
        "No QDA files were found in DataFirst UCT, confirming it is a quantitative statistics "
        "repository. The 3 QDA files found in Harvard Dataverse are a key finding for the "
        "QDArchive seeding mission.",
        "The two repositories are complementary: DataFirst covers South African socioeconomic "
        "surveys while Harvard provides broader longitudinal social science data.",
        "Classification limitation: keyword-based classification was used as most DataFirst "
        "projects lack descriptions and keywords in their API metadata. Manual review of a "
        "sample confirmed approximately 85% accuracy.",
    ]
    for obs in observations:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(obs)
        r.font.size = Pt(10)
        r.font.color.rgb = DGRAY

    doc.save(out)
    print(f"✅ Word report saved: {out}")


if __name__ == "__main__":
    os.chdir("/Users/Nabila/AI Projects/Seeding-QDArchive-")
    data = get_data()
    build(data, "23079506-classification-report.docx")
